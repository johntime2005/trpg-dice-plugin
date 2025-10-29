"""
TRPG 文档管理模块

提供文档处理、向量化存储和智能检索功能。
"""

import io
import json
import uuid
import asyncio
from typing import List, Dict, Optional, Any
from datetime import datetime

# 从nekro_agent.api.core导入Qdrant客户端
from nekro_agent.api import core
from nekro_agent.services.agent.openai import gen_openai_embeddings

# 可选依赖的导入
try:
    import docx  # python-docx for Word documents

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2  # PyPDF2 for PDF documents

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class DocumentProcessor:
    """
    文档处理器 - 支持多种格式的文档解析和向量化

    支持的格式:
    - TXT: 纯文本文件
    - PDF: PDF文档 (需要PyPDF2)
    - DOCX: Microsoft Word文档 (需要python-docx)
    """

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """从TXT文件提取文本"""
        try:
            # 尝试多种编码
            for encoding in ["utf-8", "gbk", "gb2312", "big5"]:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # 如果都失败，使用错误忽略模式
            return file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"无法解析TXT文件: {e}")

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """从PDF文件提取文本"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF处理功能不可用，请安装PyPDF2库")

        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)

            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())

            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(f"无法解析PDF文件: {e}")

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """从DOCX文件提取文本"""
        if not DOCX_AVAILABLE:
            raise ValueError("Word文档处理功能不可用，请安装python-docx库")

        try:
            docx_stream = io.BytesIO(file_content)
            doc = docx.Document(docx_stream)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(f"无法解析Word文档: {e}")

    @staticmethod
    def extract_text_by_extension(filename: str, file_content: bytes) -> str:
        """根据文件扩展名提取文本"""
        extension = filename.lower().split(".")[-1]

        if extension == "txt":
            return DocumentProcessor.extract_text_from_txt(file_content)
        elif extension == "pdf":
            return DocumentProcessor.extract_text_from_pdf(file_content)
        elif extension in ["docx", "doc"]:
            return DocumentProcessor.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"不支持的文件格式: {extension}")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 400) -> List[str]:
        """
        分块处理文本 - text-embedding-v4支持8192 token长度，可以适当增大块大小

        Args:
            text: 要分块的文本
            chunk_size: 每块的最大字符数（text-embedding-v4支持更长输入）
            overlap: 块之间的重叠字符数

        Returns:
            List[str]: 分块后的文本列表
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # 如果不是最后一块，尝试在合适的位置断开
            if end < len(text):
                # 寻找最近的句号、问号、感叹号或换行符
                for break_char in ["\n\n", "。", "！", "？", "\n", "，", "；"]:
                    break_pos = text.rfind(break_char, start, end)
                    if break_pos > start:
                        end = break_pos + len(break_char)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # 下一块的开始位置，考虑重叠
            start = max(start + chunk_size - overlap, end)

        return [chunk for chunk in chunks if chunk]


class VectorDatabaseManager:
    """
    向量数据库管理器 - 处理文档的向量化存储和检索
    """

    def __init__(self, collection_name: str = "trpg_documents"):
        self.collection_name = collection_name
        self.qdrant_client = None
        self.embedding_dim = 1536  # text-embedding-v4支持1536维度，默认使用1536
        self.document_processor = DocumentProcessor()

    async def _get_client(self):
        """获取Qdrant客户端"""
        if self.qdrant_client is None:
            self.qdrant_client = await core.get_qdrant_client()
            if not self.qdrant_client:
                raise ConnectionError("无法获取Qdrant客户端实例！请检查Agent核心配置。")
        return self.qdrant_client

    async def _ensure_collection_exists(self):
        """确保集合存在"""
        client = await self._get_client()

        try:
            # 检查集合是否存在
            collection_info = await client.get_collection(self.collection_name)
            self.embedding_dim = collection_info.config.params.vectors.size
        except Exception:
            # 集合不存在，创建新集合
            # 先生成一个测试嵌入来获取维度
            try:
                # 获取模型组配置信息
                from nekro_agent.api.core import config as core_config

                model_group_info = core_config.get_model_group_info("text-embedding")
                if not model_group_info:
                    core.logger.error(
                        "无法找到 'text-embedding' 模型组配置，使用默认维度1536"
                    )
                    self.embedding_dim = 1536
                else:
                    # 调用嵌入模型生成测试向量 - text-embedding-v3支持512, 768, 1024维度
                    test_embedding = await gen_openai_embeddings(
                        model=model_group_info.CHAT_MODEL,
                        input="test",
                        api_key=model_group_info.API_KEY,
                        base_url=model_group_info.BASE_URL,
                        dimensions=1536,  # 使用1536维度，text-embedding-v4支持
                    )
                    self.embedding_dim = len(test_embedding)
            except Exception as e:
                # 如果嵌入生成失败，使用默认维度768（text-embedding-v3支持）
                core.logger.error(f"生成测试嵌入失败: {e}，使用默认维度768")
                self.embedding_dim = 768

            await client.create_collection(
                collection_name=self.collection_name,
                vectors_config={"size": self.embedding_dim, "distance": "Cosine"},
            )

    async def store_document(
        self,
        document_id: str,
        filename: str,
        text_content: str,
        chat_key: str,
        document_type: str = "module",
    ) -> int:
        """
        存储文档到向量数据库

        Args:
            document_id: 文档唯一ID
            filename: 文件名
            text_content: 文档文本内容
            user_id: 用户ID
            chat_key: 聊天会话ID
            document_type: 文档类型 (module/rule/story/background)

        Returns:
            int: 存储的文档块数量
        """
        await self._ensure_collection_exists()
        client = await self._get_client()

        # 分割文档 - text-embedding-v4支持8192 token长度，可以适当增大块大小
        chunks = self.document_processor.chunk_text(
            text_content, chunk_size=2000, overlap=400
        )

        # 生成向量并存储每个块
        points = []
        for i, chunk in enumerate(chunks):
            # 生成嵌入向量 - 使用正确的API路径
            try:
                # 获取模型组配置信息
                from nekro_agent.api.core import config as core_config

                model_group_info = core_config.get_model_group_info("text-embedding")
                if not model_group_info:
                    core.logger.error(
                        "无法找到 'text-embedding' 模型组配置，跳过嵌入生成"
                    )
                    continue

                # 调用嵌入模型生成向量
                core.logger.debug(
                    f"调用嵌入模型: model={model_group_info.CHAT_MODEL}, base_url={model_group_info.BASE_URL}"
                )
                core.logger.debug(
                    f"API密钥: {model_group_info.API_KEY[:10]}..."
                    if model_group_info.API_KEY
                    else "API密钥为空"
                )

                try:
                    embedding = await gen_openai_embeddings(
                        model=model_group_info.CHAT_MODEL,
                        input=chunk,
                        api_key=model_group_info.API_KEY,
                        base_url=model_group_info.BASE_URL,
                        dimensions=1536,
                    )
                except Exception as embed_error:
                    core.logger.error(f"嵌入生成调用失败: {embed_error}")
                    core.logger.error(f"模型: {model_group_info.CHAT_MODEL}")
                    core.logger.error(f"基础URL: {model_group_info.BASE_URL}")
                    raise

                # 验证嵌入向量维度 - text-embedding-v4支持1536维度
                vector_dimension = len(embedding)
                if vector_dimension != 1536:
                    core.logger.warning(f"嵌入向量维度不匹配！期望: 1536, 实际: {vector_dimension}")
                else:
                    core.logger.debug(f"嵌入向量维度正确: 1536")

            except Exception as e:
                # 如果嵌入生成失败，跳过这个块
                core.logger.error(f"生成文档块嵌入失败: {e}，跳过该块")
                continue

            # 创建点数据 - 使用UUID格式作为点ID
            chunk_point_id = str(uuid.uuid4())
            point = {
                "id": chunk_point_id,
                "vector": embedding,
                "payload": {
                    "document_id": document_id,
                    "filename": filename,
                    "chunk_index": i,
                    "text": chunk,
                    "chat_key": chat_key,
                    "document_type": document_type,
                    "created_at": datetime.now().isoformat(),
                },
            }
            points.append(point)

        # 批量插入
        await client.upsert(collection_name=self.collection_name, points=points)

        return len(chunks)

    async def search_documents(
        self,
        query: str,
        chat_key: str,
        document_type: Optional[str] = None,
        limit: int = 5,
    ) -> List[Dict[str, Any]]:
        """
        搜索文档内容 - text-embedding-v3优化版本

        Args:
            query: 搜索查询
            user_id: 用户ID
            chat_key: 聊天会话ID
            document_type: 可选的文档类型过滤
            limit: 返回结果数量限制

        Returns:
            List[Dict]: 搜索结果列表
        """
        await self._ensure_collection_exists()
        client = await self._get_client()

        # 生成查询向量 - text-embedding-v4支持8192 token长度，可以处理更长查询
        try:
            # 获取模型组配置信息
            from nekro_agent.api.core import config as core_config

            model_group_info = core_config.get_model_group_info("text-embedding")
            if not model_group_info:
                # 使用默认配置生成嵌入 - text-embedding-v4支持768维度
                query_embedding = await gen_openai_embeddings(
                    model="text-embedding-v4",
                    input=query,
                    api_key="",
                    base_url="",
                    dimensions=1536,
                )
            else:
                query_embedding = await gen_openai_embeddings(
                    model=model_group_info.CHAT_MODEL,
                    input=query,
                    api_key=model_group_info.API_KEY,
                    base_url=model_group_info.BASE_URL,
                    dimensions=1536,
                )
        except Exception as e:
            # 如果嵌入生成失败，使用简单的文本匹配作为备选方案
            core.logger.warning(f"生成查询嵌入失败: {e}，使用文本匹配模式")
            # 返回空列表，让上层处理
            return []

        # 构建过滤条件 - 只根据聊天环境过滤，实现群组共享
        filter_conditions = {
            "must": [{"key": "chat_key", "match": {"value": chat_key}}]
        }

        if document_type:
            filter_conditions["must"].append(
                {"key": "document_type", "match": {"value": document_type}}
            )

        # 执行搜索
        search_result = await client.search(
            collection_name=self.collection_name,
            query_vector=query_embedding,
            query_filter=filter_conditions,
            limit=limit,
        )

        # 格式化结果
        results = []
        for hit in search_result:
            result = {
                "id": hit.id,
                "score": hit.score,
                "text": hit.payload["text"],
                "filename": hit.payload["filename"],
                "document_id": hit.payload["document_id"],
                "document_type": hit.payload["document_type"],
                "chunk_index": hit.payload["chunk_index"],
            }
            results.append(result)

        return results

    async def delete_document(self, document_id: str, chat_key: str) -> bool:
        """删除文档"""
        await self._ensure_collection_exists()
        client = await self._get_client()

        try:
            # 删除该文档的所有块 - 只根据文档ID和聊天环境，不再检查用户ID
            await client.delete(
                collection_name=self.collection_name,
                points_selector={
                    "filter": {
                        "must": [
                            {"key": "document_id", "match": {"value": document_id}},
                            {"key": "chat_key", "match": {"value": chat_key}},
                        ]
                    }
                },
            )
            return True
        except Exception:
            return False

    async def list_documents(
        self, chat_key: str, document_type: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """列出用户的所有文档"""
        await self._ensure_collection_exists()
        client = await self._get_client()

        # 构建过滤条件 - 只根据聊天环境过滤，实现群组共享
        filter_conditions = {
            "must": [
                {"key": "chat_key", "match": {"value": chat_key}},
                {
                    "key": "chunk_index",
                    "match": {"value": 0},
                },  # 只获取第一个块来代表文档
            ]
        }

        if document_type:
            filter_conditions["must"].append(
                {"key": "document_type", "match": {"value": document_type}}
            )

        # 执行搜索
        search_result = await client.scroll(
            collection_name=self.collection_name,
            scroll_filter=filter_conditions,
            limit=100,
        )

        # 格式化结果
        documents = []
        for hit in search_result[0]:  # scroll返回(points, next_page_offset)
            doc = {
                "document_id": hit.payload["document_id"],
                "filename": hit.payload["filename"],
                "document_type": hit.payload["document_type"],
                "created_at": hit.payload["created_at"],
                "preview": hit.payload["text"][:100] + "..."
                if len(hit.payload["text"]) > 100
                else hit.payload["text"],
            }
            documents.append(doc)

        return documents

    async def get_document_context(
        self, query: str, chat_key: str, max_context_length: int = 2000
    ) -> str:
        """获取与查询相关的文档上下文"""
        search_results = await self.search_documents(query, chat_key, limit=5)

        if not search_results:
            return ""

        context_parts = []
        current_length = 0

        for result in search_results:
            text = result["text"]
            filename = result["filename"]

            # 添加文档信息
            chunk_info = f"\n--- 来自文档: {filename} ---\n{text}\n"

            if current_length + len(chunk_info) > max_context_length:
                break

            context_parts.append(chunk_info)
            current_length += len(chunk_info)

        return "\n".join(context_parts)

    async def answer_question(self, question: str, chat_key: str) -> str:
        """基于文档内容回答问题 - text-embedding-v4优化版本"""
        # 搜索相关文档 - text-embedding-v4支持更长查询，提升搜索准确性
        context = await self.get_document_context(
            question, chat_key, max_context_length=3000
        )

        if not context.strip():
            return "抱歉，我在您上传的文档中没有找到相关信息来回答这个问题。"

        # 构建提示词 - text-embedding-v4提升效果，支持多语言
        prompt = f"""基于以下文档内容回答问题。如果文档中没有相关信息，请明确说明。

问题: {question}

相关文档内容:
{context}

请基于上述文档内容回答问题，如果信息不足请说明。"""

        # 这里需要调用AI模型来生成回答
        # 具体实现取决于nekro agent的AI调用接口
        try:
            from nekro_agent import core

            answer = await core.call_ai_with_prompt(prompt)
            return answer
        except Exception:
            return "抱歉，在处理您的问题时遇到了技术问题。"
